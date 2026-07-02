# -*- coding: utf-8 -*-
"""Generate dokumen Word: The Century Project — 100 achievement SakuKilat."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# (No, Nama, Syarat, Dopamin Copy, Trigger)
ROWS = [
    # A. Streak & Kebiasaan (12)
    (1,"Pecah Telur","Catat transaksi pertama","Telurnya pecah! Perjalanan dompet sehatmu dimulai.","ON_TX_SUBMIT"),
    (2,"Seminggu Tegak","Streak 7 hari","7 hari tanpa absen. Disiplinmu mulai kelihatan.","ON_TX_SUBMIT"),
    (3,"Sebulan Penuh","Streak 30 hari","Sebulan nonstop! Ini bukan kebetulan, ini kebiasaan.","ON_TX_SUBMIT"),
    (4,"Si Kepala Batu","Streak 100 hari","100 hari?! Kepala batu finansial sejati.","ON_TX_SUBMIT"),
    (5,"Setahun Setia","Streak 365 hari","Setahun penuh. Kamu legenda, titik.","ON_TX_SUBMIT"),
    (6,"Konsisten Sejati","90 hari aktif total mencatat","90 hari penuh jejak. Konsistensi level dewa.","ON_APP_MOUNT"),
    (7,"Pejuang Subuh","Catat sebelum jam 6 pagi 7 kali","Belum melek penuh tapi dompet udah dicatat. Salut.","ON_TX_SUBMIT"),
    (8,"Anak Weekend","Catat di Sabtu dan Minggu","Libur boleh, catat jalan terus.","ON_TX_SUBMIT"),
    (9,"Penunggu Malam","Catat di malam hari (18-23)","Tutup hari dengan catatan. Tidurmu lebih tenang.","ON_TX_SUBMIT"),
    (10,"Dua Minggu Beruntun","Rekor streak 14 hari","Dua minggu nonstop. Momentummu ngebut.","ON_TX_SUBMIT"),
    (11,"Bangkit Lagi","Streak 7 hari setelah sempat putus","Jatuh bukan akhir. Kamu bangkit dan lari lagi.","ON_TX_SUBMIT"),
    (12,"Nyawa Utuh","Catat hari ini dengan 5 nyawa penuh","5 nyawa aman. Rajinmu nggak main-main.","ON_TX_SUBMIT"),
    # B. Pencapaian & Volume (15)
    (13,"Setengah Ratus","50 transaksi tercatat","50 catatan! Dompetmu makin terbaca.","ON_TX_SUBMIT"),
    (14,"Lima Ratus Jejak","500 transaksi","500 jejak keuangan. Arsiparis beneran.","ON_TX_SUBMIT"),
    (15,"Dua Ribu Lima Ratus","2.500 transaksi","2.500 catatan. Ini sih hobi, bukan tugas lagi.","ON_TX_SUBMIT"),
    (16,"Maha Pencatat","10.000 transaksi","10.000! Tugu peringatan layak dibangun untukmu.","ON_TX_SUBMIT"),
    (17,"Juta Pertama Tercatat","Total pengeluaran tercatat tembus Rp1 juta","Sejuta rupiah terlacak. Nggak ada yang lolos.","ON_TX_SUBMIT"),
    (18,"Sultan Pencatat","Total pengeluaran tercatat tembus Rp100 juta","100 juta lewat tanganmu, semua tercatat rapi.","ON_TX_SUBMIT"),
    (19,"Dompet Lengkap","Punya 5 saku","5 saku aktif. Manajer dompet profesional.","ON_ROUTE_CHANGE"),
    (20,"Kolektor Kategori","Pakai 10 kategori berbeda","10 warna pengeluaran. Hidupmu berwarna (dan tercatat).","ON_TX_SUBMIT"),
    (21,"Goal Perdana","Buat goal tabungan pertama","Mimpi pertama dipasang. Ayo dikejar!","ON_ROUTE_CHANGE"),
    (22,"Mimpi Terwujud","Capai 1 goal","Target tembus! Rasanya beda kan kalau direncanakan.","ON_TX_SUBMIT"),
    (23,"Pemburu Lima Mimpi","Capai 5 goal","5 mimpi terwujud. Kamu mesin pewujud target.","ON_TX_SUBMIT"),
    (24,"Backup Perdana","Backup/ekspor pertama","Data diamankan. Tidur lebih nyenyak.","ON_ROUTE_CHANGE"),
    (25,"Tukang Arsip","Backup 10 kali","10 backup. Paranoid sehat, kami suka.","ON_ROUTE_CHANGE"),
    (26,"Migrasi Sukses","Impor data pertama kali","Pindah rumah lancar. Data lama selamat sampai tujuan.","ON_APP_MOUNT"),
    (27,"Veteran Data","Pernah impor & 10 kali backup","Pindahan + rajin backup. Anti kehilangan sejati.","ON_ROUTE_CHANGE"),
    # C. Disiplin Anggaran (20)
    (28,"Set Budget Pertama","Tetapkan budget bulanan","Pagar sudah dipasang. Sekarang tinggal jaga.","ON_ROUTE_CHANGE"),
    (29,"Penabung 30%","Rasio tabungan tembus 30% sebulan","30% disisihkan. Masa depanmu berterima kasih.","ON_TX_SUBMIT"),
    (30,"Penabung 50%","Rasio tabungan tembus 50% sebulan","Separuh disimpan?! Disiplin baja.","ON_TX_SUBMIT"),
    (31,"Surplus Bulanan","Pemasukan > pengeluaran sebulan","Plus, bukan minus. Bulan ini kamu menang.","ON_CRON_MIDNIGHT"),
    (32,"Bebas Utang","Semua saku bersaldo positif","Nggak ada yang minus. Lega tanpa beban.","ON_TX_SUBMIT"),
    (33,"Rajin Nabung","Nabung ke tabungan 30 hari berbeda","Sebulan nabung tiap hari. Tetes demi tetes jadi danau.","ON_TX_SUBMIT"),
    (34,"Celengan Tebal","Saldo Tabungan tembus Rp5 juta","Celengan gemuk. Bunyinya udah berat.","ON_TX_SUBMIT"),
    (35,"Multi Cuan","Pemasukan dari 3 kategori berbeda","Banyak keran cuan. Nggak gantung satu sumber.","ON_TX_SUBMIT"),
    (36,"Mindful Spender","Buka Rekapan untuk evaluasi","Evaluasi dulu, belanja kemudian. Bijak.","ON_ROUTE_CHANGE"),
    (37,"Detektif Duit","Buka Rekapan di 30 hari berbeda","Rajin investigasi dompet sendiri. Nggak ada misteri.","ON_ROUTE_CHANGE"),
    (38,"Tukang Rapi","Semua transaksi punya kategori (min. 20 tx)","Nggak ada yang nyasar ke 'lainnya'. Rapi total.","ON_TX_SUBMIT"),
    (39,"Si Bijak","Budget, goal, & backup aktif bersamaan","Pagar, mimpi, dan cadangan lengkap. Paket komplit.","ON_ROUTE_CHANGE"),
    (40,"Tutup Bulan Aman","Akhiri bulan di bawah budget","Bulan ditutup dengan senyum. Budget terjaga.","ON_CRON_MIDNIGHT"),
    (41,"Hemat Trilogi","3 bulan beruntun di bawah budget","Tiga bulan irit berturut. Ini bukan keberuntungan.","ON_CRON_MIDNIGHT"),
    (42,"Master Anggaran","12 bulan keuangan sehat","Setahun anggaran terkendali. Sensei budgeting.","ON_CRON_MIDNIGHT"),
    (43,"Di Bawah Separuh","Pakai <50% budget sebulan","Setengah budget aja cukup. Hemat tingkat lanjut.","ON_CRON_MIDNIGHT"),
    (44,"Irit Maksimal","Pakai <25% budget sebulan","Cuma seperempat budget?! Ajarin dong.","ON_CRON_MIDNIGHT"),
    (45,"Minggu Hijau","4 minggu tanpa over jatah","Empat minggu hijau berturut. Stabil banget.","ON_CRON_MIDNIGHT"),
    (46,"Anti Bocor","Sebulan tanpa satu hari over jatah","Nol kebocoran sebulan. Dompet kedap air.","ON_CRON_MIDNIGHT"),
    (47,"Budget Naik Kelas","Revisi budget setelah evaluasi","Budget di-upgrade berdasarkan data. Makin matang.","ON_ROUTE_CHANGE"),
    # D. Psikologi & Zen (18)
    (48,"Zen Master","Aktifkan Zen Mode","Angka disembunyikan. Pikiran lebih damai.","ON_ROUTE_CHANGE"),
    (49,"Filosof Dompet","Pakai Zen Mode 30 hari","Sebulan dalam ketenangan. Uang bukan tuanmu.","ON_CRON_MIDNIGHT"),
    (50,"Hari Tanpa Jajan","1 hari penuh tanpa pengeluaran","Seharian nol jajan. Dompetmu istirahat.","ON_CRON_MIDNIGHT"),
    (51,"Puasa Belanja","3 hari beruntun tanpa pengeluaran","3 hari puasa belanja. Tahan godaan, naik level.","ON_CRON_MIDNIGHT"),
    (52,"Akhir Pekan Hemat","Sabtu-Minggu tanpa pengeluaran","Weekend nol jajan. Healing nggak harus mahal.","ON_CRON_MIDNIGHT"),
    (53,"Lebih Irit dari Lalu","Pengeluaran mingguan turun","Minggu ini lebih hemat. Grafik turun, hati senang.","ON_CRON_MIDNIGHT"),
    (54,"Refleksi Tenang","Buka app & lihat data tanpa belanja","Sekadar merenungi angka. Sadar diri itu kaya.","ON_APP_MOUNT"),
    (55,"Pembaca Data","Buka Rekapan di 10 hari berbeda","Rajin baca laporan sendiri. Nggak buta arah.","ON_ROUTE_CHANGE"),
    (56,"Frugal Sejati","Tabungan >50% selama 3 bulan","Tiga bulan hemat ekstrem. Mindset kaya beneran.","ON_CRON_MIDNIGHT"),
    (57,"Napas Panjang","Savings rate positif 6 bulan","Setengah tahun selalu nyisihkan. Napas finansialmu panjang.","ON_CRON_MIDNIGHT"),
    (58,"Kepala Dingin","Seminggu tanpa belanja impulsif >Rp500rb","Nggak ada checkout panas. Kepala tetap dingin.","ON_CRON_MIDNIGHT"),
    (59,"Anti FOMO","Seminggu tanpa pengeluaran hiburan","Skip hiburan seminggu. FOMO kalah sama logika.","ON_CRON_MIDNIGHT"),
    (60,"Sadar Diri","Catat dengan jujur 30 hari aktif","Sebulan jujur sama dompet sendiri. Itu langka.","ON_APP_MOUNT"),
    (61,"Si Evaluator","Buka tampilan Tren di Rekapan","Lihat tren, ambil pelajaran. Otak finansial nyala.","ON_ROUTE_CHANGE"),
    (62,"Perencana Ulung","Buat goal dengan tenggat waktu","Mimpi dengan deadline. Itu rencana, bukan angan.","ON_ROUTE_CHANGE"),
    (63,"Hidup Seimbang","Pemasukan & pengeluaran tercatat di bulan sama","Masuk dan keluar seimbang tercatat. Gambaran utuh.","ON_TX_SUBMIT"),
    (64,"Tepat Deadline","Capai goal sebelum/tepat tenggat","Target kelar tepat waktu. Perencanaan jempolan.","ON_TX_SUBMIT"),
    (65,"Mahir Mengelola","3+ saku dan minimal 1 goal aktif","Banyak dompet, terarah ke tujuan. Pengelola handal.","ON_ROUTE_CHANGE"),
    # E. Lore & Easter Egg (35)
    (66,"Gaji Numpang Lewat","Income ≥Rp3jt, lalu tergerus >90% dalam 5 hari","Gaji cuma mampir say hi. Sabar, akhir bulan masih jauh.","ON_TX_SUBMIT"),
    (67,"Survivor Tanggal Tua","5 hari (tgl 20-25) pengeluaran harian <Rp20rb","Bertahan di tanggal tua dengan elegan. Hidup keras, kamu lebih keras.","ON_CRON_MIDNIGHT"),
    (68,"Pawang Parkir","10 transaksi tunai tepat Rp2.000","Receh parkir terlacak semua. Nggak ada yang lolos, pak.","ON_TX_SUBMIT"),
    (69,"Budak Paylater","Pengeluaran pertama via metode 'Paylater'","Beli sekarang, nangis nanti. Tercatat ya, jangan lupa.","ON_TX_SUBMIT"),
    (70,"Gak Jadi Beli","Pakai Urungkan 3 kali","Maju mundur cantik. Akhirnya nggak jadi beli juga, hemat!","ON_TX_SUBMIT"),
    (71,"Korban Diskon","Belanja >Rp300rb di tanggal kembar","Diskon emang jebakan. Tapi tercatat, jadi nggak sepenuhnya kalah.","ON_TX_SUBMIT"),
    (72,"Racun Checkout Malam","Belanja online jam 00-03","Jempol gatel tengah malam. Besok pagi baru nyesel.","ON_TX_SUBMIT"),
    (73,"Sultan Sehari","Satu transaksi >Rp10 juta","Sekali transaksi, gaji orang sebulan. Hormat, bos.","ON_TX_SUBMIT"),
    (74,"Receh Hunter","Catat transaksi <Rp1.000","Recehan pun nggak luput. Detail banget kamu.","ON_TX_SUBMIT"),
    (75,"Caffeine Dependent","Catat 'kopi' 7 kali","Tujuh kali ngopi. Dompet & jantung sama-sama deg-degan.","ON_TX_SUBMIT"),
    (76,"Korban Boba","Catat transaksi boba/milk tea","Boba lagi, boba lagi. Manisnya nempel di pengeluaran.","ON_TX_SUBMIT"),
    (77,"Ojol Setia","10 transaksi ojek online","Mitra setia ojol. Abang driver berterima kasih.","ON_TX_SUBMIT"),
    (78,"Anak Minimarket","10 transaksi di minimarket","Mampir 'cuma beli air', keluar bawa kresek. Klasik.","ON_TX_SUBMIT"),
    (79,"Korban Ongkir","Catat 5 transaksi 'ongkir'","Barang Rp10rb, ongkir Rp20rb. Logika belanja online.","ON_TX_SUBMIT"),
    (80,"Dompet Tipis","Total saldo semua saku <Rp50rb","Tinggal segini? Tarik napas, akhir bulan ujian sesungguhnya.","ON_TX_SUBMIT"),
    (81,"Tajir Mendadak","Pemasukan >Rp5 juta sekaligus","Dari mana nih durian runtuh? Selamat menikmati (sebentar).","ON_TX_SUBMIT"),
    (82,"THR Cair!","Catat pemasukan THR/bonus/hadiah","THR turun! Tahan, jangan langsung ludes ya.","ON_TX_SUBMIT"),
    (83,"Gajian!","Catat pemasukan kategori gaji","Saldo hijau lagi! Tarik napas, ini cuma titipan tagihan.","ON_TX_SUBMIT"),
    (84,"Tekor Awal Bulan","Pengeluaran besar di tanggal 1-5","Baru awal bulan udah ngebut. Hati-hati, finish line jauh.","ON_TX_SUBMIT"),
    (85,"Anak Padang","Catat 'padang' 5 kali","Rendang lover sejati. Lauknya boleh, dompetnya dijaga.","ON_TX_SUBMIT"),
    (86,"Patungan Pro","Pakai fitur bagi (split bill)","Bayar bareng, hemat bareng. Temen-temen sayang kamu.","ON_TX_SUBMIT"),
    (87,"Tukang Geser","10 kali pindah uang antar saku","Geser sana geser sini. Bendahara grup ya?","ON_TX_SUBMIT"),
    (88,"Nabung Tengah Malam","Simpan uang jam 00-04","Insaf tengah malam, langsung nabung. Hidayah finansial.","ON_TX_SUBMIT"),
    (89,"Begadang Finansial","Catat transaksi jam 00-04","Mata panda, tapi dompet tetap tercatat. Respect.","ON_TX_SUBMIT"),
    (90,"Voice Note Master","Pakai input suara 10 kali","Ngomong doang, langsung tercatat. Generasi rebahan.","ON_TX_SUBMIT"),
    (91,"Ketik Kilat","Pakai singkatan jt/k/rb saat mencatat","Ngetik 5jt bukan 5000000. Time is money, literally.","ON_TX_SUBMIT"),
    (92,"Si Telat Catat","Catat transaksi 3+ hari lalu","Telat tapi tetap dicatat. Mending telat daripada lupa.","ON_TX_SUBMIT"),
    (93,"Tukang Edit","Edit 10 transaksi","Perfeksionis dompet. Harus pas sampai koma terakhir.","ON_TX_SUBMIT"),
    (94,"Pembaca Setia","Buka Buku Panduan","Baca dulu sebelum nanya. Kamu user idaman.","ON_ROUTE_CHANGE"),
    (95,"Kepo Fitur","Buka semua tab","Diubek-ubek semua fiturnya. Rasa penasaran tingkat dewa.","ON_ROUTE_CHANGE"),
    (96,"Ganti Wajah","Ubah foto profil","Tampil beda. Dompet rapi, profil juga harus kece.","ON_ROUTE_CHANGE"),
    (97,"Kolektor Saku","Punya 8 saku berbeda","Dompet bercabang ke mana-mana. Sultan multi-rekening.","ON_ROUTE_CHANGE"),
    (98,"Si Kreatif","Buat kategori kustom sendiri","Kategori bawaan kurang? Bikin sendiri, bos. Merdeka!","ON_ROUTE_CHANGE"),
    (99,"Komplit Sehari","Pemasukan, pengeluaran, & pindah uang dalam sehari","Triple combo dalam sehari. Aktivitas dompet padat merayap.","ON_CRON_MIDNIGHT"),
    (100,"Khatam SakuKilat","Buka semua tab, panduan, & raih 50 lencana","Tamat sudah! Kamu menguasai SakuKilat luar dalam.","ON_ROUTE_CHANGE"),
]

CYAN = RGBColor(0x0E, 0x76, 0x90)
GRAY = RGBColor(0x60, 0x6A, 0x7B)

doc = Document()
title = doc.add_heading("The Century Project — 100 Achievement SakuKilat", level=0)
sub = doc.add_paragraph("100 lencana kurasi: anti-filler, kaya lore Indonesia, dan ber-trigger event "
                        "untuk efisiensi CPU. Semua dihitung dari data lokal pengguna.")
sub.runs[0].italic = True
sub.runs[0].font.color.rgb = GRAY

cap = doc.add_paragraph("Distribusi: Streak & Kebiasaan 12 | Pencapaian & Volume 15 | "
                        "Disiplin Anggaran 20 | Psikologi & Zen 18 | Lore & Easter Egg 35.")
cap.runs[0].bold = True
doc.add_paragraph("")

table = doc.add_table(rows=1, cols=5)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, h in enumerate(["No", "Nama Lencana", "Syarat Nyata", "Dopamin Copy (Pop-up)", "Eval_Trigger"]):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True

for no, nama, syarat, copy, trig in ROWS:
    row = table.add_row().cells
    row[0].text = str(no)
    row[1].text = nama
    row[2].text = syarat
    row[3].text = copy
    row[4].text = trig

doc.add_paragraph("")
foot = doc.add_paragraph("Eval_Trigger menentukan KAPAN lencana dicek: ON_TX_SUBMIT (saat simpan transaksi), "
                         "ON_APP_MOUNT (saat app dibuka), ON_ROUTE_CHANGE (saat pindah tab), "
                         "ON_CRON_MIDNIGHT (dicek senyap tengah malam). Tujuannya agar app tidak menghitung "
                         "100 fungsi tiap ketukan keyboard.")
foot.runs[0].font.size = Pt(9)
foot.runs[0].font.color.rgb = GRAY

out = "docs/SakuKilat-100-Achievements.docx"
doc.save(out)
print(f"SAVED {out} | total = {len(ROWS)}")
