# -*- coding: utf-8 -*-
"""Panduan Lengkap SakuKilat: keunggulan + cara pakai + 100 achievement & cara dapat."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CYAN = RGBColor(0x0E, 0x76, 0x90)
GRAY = RGBColor(0x60, 0x6A, 0x7B)
GREEN = RGBColor(0x1E, 0x8E, 0x5A)

doc = Document()

# ── Sampul ──
t = doc.add_heading("Panduan Lengkap SakuKilat", level=0)
s = doc.add_paragraph("Pencatat keuangan lokal-first berbahasa Indonesia. "
                      "Catat secepat kilat, pahami uangmu, dan kumpulkan lencana.")
s.runs[0].italic = True
s.runs[0].font.color.rgb = GRAY
doc.add_paragraph("")

# ── Bagian 1: Keunggulan ──
doc.add_heading("1. Keunggulan Utama", level=1).runs[0].font.color.rgb = CYAN
keunggulan = [
    ("Catat secepat ngetik chat", "Tulis bahasa natural seperti 'kopi 18k gopay'. Nominal, kategori, dan saku ditebak otomatis."),
    ("100% lokal & offline", "Semua data tersimpan di perangkatmu. Tanpa login, tanpa server, tetap jalan tanpa internet."),
    ("Input suara", "Tap tombol mic lalu ucapkan transaksimu dalam Bahasa Indonesia."),
    ("Rekap kaya", "History, kalender intensitas harian, dan grafik tren dalam satu tab."),
    ("Budget pintar", "Batas aman jajan harian yang menyesuaikan otomatis sesuai sisa anggaran."),
    ("Gamifikasi sehat", "Streak harian, 5 nyawa, dan 100 lencana membuat mencatat jadi kebiasaan."),
    ("Backup & migrasi", "Ekspor/impor JSON & CSV. Pindah perangkat tanpa kehilangan data."),
]
for judul, isi in keunggulan:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(judul + ": "); r.bold = True
    p.add_run(isi)
doc.add_paragraph("")

# ── Bagian 2: Cara Pakai per Fitur ──
doc.add_heading("2. Cara Memakai Tiap Fitur", level=1).runs[0].font.color.rgb = CYAN
fitur = [
    ("Smart Tracker (catat cepat)",
     "Rumus: apa + berapa + pakai apa. Contoh: 'kopi 18k gopay', 'gaji 5jt bca', 'bensin 50rb tunai'. "
     "Singkatan: k/rb/ribu = ribuan, jt/juta = jutaan. Bisa split bill ('bagi 4') dan tanggal ('kemarin', 'tgl 5'). "
     "Tekan Enter untuk menyimpan."),
    ("Catat Manual & Suara",
     "Tombol mic untuk input suara. Tombol slider membuka form manual untuk memilih tipe, dompet, kategori, "
     "tanggal, dan deskripsi sendiri — cocok untuk transaksi rumit atau lampau."),
    ("Saku & Budget",
     "Di tab Saku, kelola dompet (cash/bank/e-wallet) dan isi saldo awal. Atur Budget Bulanan agar Beranda "
     "bisa mengingatkan sisa jatah harian & mingguan."),
    ("Pindah, Simpan & Goal",
     "Geser uang antar saku tanpa dihitung pengeluaran (bisa juga 'pindah 100k bca ke gopay'). "
     "Buat Goal Tabungan dan tambahkan kontribusi sedikit demi sedikit."),
    ("Rekapan",
     "Tiga tampilan: History (daftar transaksi, filter periode), Kalender (intensitas harian), "
     "dan Tren (grafik + alokasi kategori). Ketuk satu hari di kalender untuk rinciannya."),
    ("Metode Bayar & Kategori",
     "Tambah metode/kategori kustom lengkap dengan keyword agar parser mengenali istilah khasmu. "
     "Kategori bisa punya sub-kategori."),
    ("Streak, Nyawa & Trofi",
     "Catat tiap hari untuk menambah streak. Ada 5 nyawa; satu pecah tiap hari absen, terisi lagi saat kamu catat. "
     "Etalase Trofi mengumpulkan lencana atas kebiasaan baik."),
    ("Notifikasi",
     "Tap lonceng di pojok kanan atas Beranda untuk pengingat: streak hampir putus, budget menipis, goal hampir tercapai."),
    ("Data, Backup & Privasi",
     "Semua data di perangkat ini saja. Backup JSON (cadangan penuh), Ekspor CSV (buka di Excel/Sheets), dan Impor. "
     "Saat impor dari aplikasi lain, saku & kategori baru dibuat otomatis."),
]
for judul, isi in fitur:
    h = doc.add_heading(judul, level=2)
    doc.add_paragraph(isi)
doc.add_paragraph("")

# ── Bagian 3: 100 Achievement & Cara Mendapatkan ──
doc.add_heading("3. 100 Lencana & Cara Mendapatkannya", level=1).runs[0].font.color.rgb = CYAN
doc.add_paragraph("Lencana terbuka otomatis saat syaratnya terpenuhi — termasuk langsung setelah mengimpor data lama. "
                  "Berikut daftar lengkap beserta cara pastinya.")

# (Grup, [(no, nama, cara_dapat, dopamin)])
GROUPS = [
 ("A. Streak & Kebiasaan", [
  (1,"Pecah Telur","Catat transaksi pertamamu.","Telurnya pecah! Perjalanan dompet sehatmu dimulai."),
  (2,"Seminggu Tegak","Catat transaksi 7 hari beruntun tanpa absen.","7 hari tanpa absen. Disiplinmu mulai kelihatan."),
  (3,"Sebulan Penuh","Catat 30 hari beruntun.","Sebulan nonstop! Ini bukan kebetulan, ini kebiasaan."),
  (4,"Si Kepala Batu","Catat 100 hari beruntun.","100 hari?! Kepala batu finansial sejati."),
  (5,"Setahun Setia","Catat 365 hari beruntun.","Setahun penuh. Kamu legenda, titik."),
  (6,"Konsisten Sejati","Kumpulkan 90 hari aktif mencatat (total, tidak harus beruntun).","90 hari penuh jejak. Konsistensi level dewa."),
  (7,"Pejuang Subuh","Catat transaksi sebelum jam 6 pagi sebanyak 7 kali.","Belum melek penuh tapi dompet udah dicatat. Salut."),
  (8,"Anak Weekend","Catat transaksi di hari Sabtu dan Minggu.","Libur boleh, catat jalan terus."),
  (9,"Penunggu Malam","Catat transaksi di malam hari (jam 18.00-23.00).","Tutup hari dengan catatan. Tidurmu lebih tenang."),
  (10,"Dua Minggu Beruntun","Capai rekor streak 14 hari.","Dua minggu nonstop. Momentummu ngebut."),
  (11,"Bangkit Lagi","Bangun lagi streak 7 hari setelah streak sebelumnya sempat putus.","Jatuh bukan akhir. Kamu bangkit dan lari lagi."),
  (12,"Nyawa Utuh","Catat hari ini saat semua 5 nyawa masih penuh.","5 nyawa aman. Rajinmu nggak main-main."),
 ]),
 ("B. Pencapaian & Volume", [
  (13,"Setengah Ratus","Kumpulkan total 50 transaksi.","50 catatan! Dompetmu makin terbaca."),
  (14,"Lima Ratus Jejak","Kumpulkan total 500 transaksi.","500 jejak keuangan. Arsiparis beneran."),
  (15,"Dua Ribu Lima Ratus","Kumpulkan total 2.500 transaksi.","2.500 catatan. Ini sih hobi, bukan tugas lagi."),
  (16,"Maha Pencatat","Kumpulkan total 10.000 transaksi.","10.000! Tugu peringatan layak dibangun untukmu."),
  (17,"Juta Pertama Tercatat","Total pengeluaran yang tercatat tembus Rp1.000.000.","Sejuta rupiah terlacak. Nggak ada yang lolos."),
  (18,"Sultan Pencatat","Total pengeluaran tercatat tembus Rp100.000.000.","100 juta lewat tanganmu, semua tercatat rapi."),
  (19,"Dompet Lengkap","Miliki 5 saku.","5 saku aktif. Manajer dompet profesional."),
  (20,"Kolektor Kategori","Pakai 10 kategori berbeda pada transaksimu.","10 warna pengeluaran. Hidupmu berwarna (dan tercatat)."),
  (21,"Goal Perdana","Buat goal tabungan pertama.","Mimpi pertama dipasang. Ayo dikejar!"),
  (22,"Mimpi Terwujud","Capai (penuhi target) 1 goal tabungan.","Target tembus! Rasanya beda kan kalau direncanakan."),
  (23,"Pemburu Lima Mimpi","Capai 5 goal tabungan.","5 mimpi terwujud. Kamu mesin pewujud target."),
  (24,"Backup Perdana","Lakukan backup atau ekspor data pertama kali.","Data diamankan. Tidur lebih nyenyak."),
  (25,"Tukang Arsip","Lakukan backup sebanyak 10 kali.","10 backup. Paranoid sehat, kami suka."),
  (26,"Migrasi Sukses","Impor data pertama kali.","Pindah rumah lancar. Data lama selamat sampai tujuan."),
  (27,"Veteran Data","Pernah impor data DAN sudah backup 10 kali.","Pindahan + rajin backup. Anti kehilangan sejati."),
 ]),
 ("C. Disiplin Anggaran", [
  (28,"Set Budget Pertama","Tetapkan budget bulanan di tab Saku.","Pagar sudah dipasang. Sekarang tinggal jaga."),
  (29,"Penabung 30%","Capai rasio tabungan (pemasukan dikurangi pengeluaran) 30% dalam sebulan.","30% disisihkan. Masa depanmu berterima kasih."),
  (30,"Penabung 50%","Capai rasio tabungan 50% dalam sebulan.","Separuh disimpan?! Disiplin baja."),
  (31,"Surplus Bulanan","Akhiri bulan dengan pemasukan lebih besar dari pengeluaran.","Plus, bukan minus. Bulan ini kamu menang."),
  (32,"Bebas Utang","Pastikan tidak ada saku yang bersaldo minus.","Nggak ada yang minus. Lega tanpa beban."),
  (33,"Rajin Nabung","Sumbang ke tabungan di 30 hari yang berbeda.","Sebulan nabung tiap hari. Tetes demi tetes jadi danau."),
  (34,"Celengan Tebal","Capai saldo saku Tabungan Rp5.000.000.","Celengan gemuk. Bunyinya udah berat."),
  (35,"Multi Cuan","Catat pemasukan dari 3 kategori berbeda.","Banyak keran cuan. Nggak gantung satu sumber."),
  (36,"Mindful Spender","Buka tab Rekapan untuk evaluasi.","Evaluasi dulu, belanja kemudian. Bijak."),
  (37,"Detektif Duit","Buka Rekapan di 30 hari yang berbeda.","Rajin investigasi dompet sendiri. Nggak ada misteri."),
  (38,"Tukang Rapi","Pastikan semua transaksi (minimal 20) punya kategori, bukan 'lainnya'.","Nggak ada yang nyasar ke 'lainnya'. Rapi total."),
  (39,"Si Bijak","Aktifkan budget, punya goal, dan sudah backup — semua sekaligus.","Pagar, mimpi, dan cadangan lengkap. Paket komplit."),
  (40,"Tutup Bulan Aman","Akhiri satu bulan dengan total pengeluaran di bawah budget.","Bulan ditutup dengan senyum. Budget terjaga."),
  (41,"Hemat Trilogi","Tutup 3 bulan beruntun di bawah budget.","Tiga bulan irit berturut. Ini bukan keberuntungan."),
  (42,"Master Anggaran","Jaga keuangan sehat selama 12 bulan.","Setahun anggaran terkendali. Sensei budgeting."),
  (43,"Di Bawah Separuh","Habiskan kurang dari 50% budget dalam sebulan.","Setengah budget aja cukup. Hemat tingkat lanjut."),
  (44,"Irit Maksimal","Habiskan kurang dari 25% budget dalam sebulan.","Cuma seperempat budget?! Ajarin dong."),
  (45,"Minggu Hijau","Lalui 4 minggu tanpa melewati jatah.","Empat minggu hijau berturut. Stabil banget."),
  (46,"Anti Bocor","Lalui sebulan tanpa satu pun hari melewati jatah harian.","Nol kebocoran sebulan. Dompet kedap air."),
  (47,"Budget Naik Kelas","Revisi/ubah nilai budget setelah evaluasi.","Budget di-upgrade berdasarkan data. Makin matang."),
 ]),
 ("D. Psikologi & Zen", [
  (48,"Zen Master","Aktifkan Zen Mode (sembunyikan angka) di Beranda.","Angka disembunyikan. Pikiran lebih damai."),
  (49,"Filosof Dompet","Pertahankan kebiasaan Zen Mode selama 30 hari.","Sebulan dalam ketenangan. Uang bukan tuanmu."),
  (50,"Hari Tanpa Jajan","Lewati satu hari penuh tanpa pengeluaran (tetap buka app).","Seharian nol jajan. Dompetmu istirahat."),
  (51,"Puasa Belanja","Lalui 3 hari beruntun tanpa pengeluaran sama sekali.","3 hari puasa belanja. Tahan godaan, naik level."),
  (52,"Akhir Pekan Hemat","Lalui Sabtu-Minggu tanpa pengeluaran.","Weekend nol jajan. Healing nggak harus mahal."),
  (53,"Lebih Irit dari Lalu","Turunkan total pengeluaran mingguan dari minggu sebelumnya.","Minggu ini lebih hemat. Grafik turun, hati senang."),
  (54,"Refleksi Tenang","Buka aplikasi dan amati datamu tanpa belanja (punya min. 5 transaksi).","Sekadar merenungi angka. Sadar diri itu kaya."),
  (55,"Pembaca Data","Buka Rekapan di 10 hari yang berbeda.","Rajin baca laporan sendiri. Nggak buta arah."),
  (56,"Frugal Sejati","Pertahankan rasio tabungan di atas 50% selama 3 bulan.","Tiga bulan hemat ekstrem. Mindset kaya beneran."),
  (57,"Napas Panjang","Pertahankan savings rate positif selama 6 bulan.","Setengah tahun selalu nyisihkan. Napas finansialmu panjang."),
  (58,"Kepala Dingin","Lalui seminggu tanpa belanja impulsif di atas Rp500.000.","Nggak ada checkout panas. Kepala tetap dingin."),
  (59,"Anti FOMO","Lalui seminggu tanpa pengeluaran kategori hiburan.","Skip hiburan seminggu. FOMO kalah sama logika."),
  (60,"Sadar Diri","Catat dengan jujur selama 30 hari aktif.","Sebulan jujur sama dompet sendiri. Itu langka."),
  (61,"Si Evaluator","Buka tampilan Tren di Rekapan.","Lihat tren, ambil pelajaran. Otak finansial nyala."),
  (62,"Perencana Ulung","Buat goal tabungan yang punya tenggat waktu.","Mimpi dengan deadline. Itu rencana, bukan angan."),
  (63,"Hidup Seimbang","Punya pemasukan dan pengeluaran tercatat di bulan yang sama.","Masuk dan keluar seimbang tercatat. Gambaran utuh."),
  (64,"Tepat Deadline","Capai goal tepat pada atau sebelum tenggatnya.","Target kelar tepat waktu. Perencanaan jempolan."),
  (65,"Mahir Mengelola","Miliki 3 saku atau lebih dan minimal 1 goal aktif.","Banyak dompet, terarah ke tujuan. Pengelola handal."),
 ]),
 ("E. Lore & Easter Egg", [
  (66,"Gaji Numpang Lewat","Catat pemasukan besar (>=Rp3jt), lalu pengeluaran menggerus lebih dari 90% income itu dalam 5 hari.","Gaji cuma mampir say hi. Sabar, akhir bulan masih jauh."),
  (67,"Survivor Tanggal Tua","Lalui 5 hari (tanggal 20-25) dengan pengeluaran harian di bawah Rp20.000.","Bertahan di tanggal tua dengan elegan. Hidup keras, kamu lebih keras."),
  (68,"Pawang Parkir","Catat 10 transaksi tunai senilai tepat Rp2.000.","Receh parkir terlacak semua. Nggak ada yang lolos, pak."),
  (69,"Budak Paylater","Catat pengeluaran pertama memakai metode bayar bernama 'Paylater' (buat metode kustomnya dulu).","Beli sekarang, nangis nanti. Tercatat ya, jangan lupa."),
  (70,"Gak Jadi Beli","Pakai tombol Urungkan sebanyak 3 kali.","Maju mundur cantik. Akhirnya nggak jadi beli juga, hemat!"),
  (71,"Korban Diskon","Belanja di atas Rp300.000 pada tanggal kembar (mis. 12 Desember).","Diskon emang jebakan. Tapi tercatat, jadi nggak sepenuhnya kalah."),
  (72,"Racun Checkout Malam","Catat belanja/hiburan antara jam 00.00-03.00.","Jempol gatel tengah malam. Besok pagi baru nyesel."),
  (73,"Sultan Sehari","Catat satu transaksi di atas Rp10.000.000.","Sekali transaksi, gaji orang sebulan. Hormat, bos."),
  (74,"Receh Hunter","Catat transaksi bernilai di bawah Rp1.000.","Recehan pun nggak luput. Detail banget kamu."),
  (75,"Caffeine Dependent","Catat transaksi mengandung kata 'kopi' sebanyak 7 kali.","Tujuh kali ngopi. Dompet & jantung sama-sama deg-degan."),
  (76,"Korban Boba","Catat transaksi mengandung 'boba' atau 'milk tea'.","Boba lagi, boba lagi. Manisnya nempel di pengeluaran."),
  (77,"Ojol Setia","Catat 10 transaksi ojek online (gojek/grab/ojol).","Mitra setia ojol. Abang driver berterima kasih."),
  (78,"Anak Minimarket","Catat 10 transaksi di minimarket (indomaret/alfamart).","Mampir 'cuma beli air', keluar bawa kresek. Klasik."),
  (79,"Korban Ongkir","Catat 5 transaksi mengandung kata 'ongkir'.","Barang Rp10rb, ongkir Rp20rb. Logika belanja online."),
  (80,"Dompet Tipis","Buat total saldo semua saku turun di bawah Rp50.000.","Tinggal segini? Tarik napas, akhir bulan ujian sesungguhnya."),
  (81,"Tajir Mendadak","Catat satu pemasukan di atas Rp5.000.000.","Dari mana nih durian runtuh? Selamat menikmati (sebentar)."),
  (82,"THR Cair!","Catat pemasukan kategori hadiah, atau deskripsi berisi 'THR'/'bonus'.","THR turun! Tahan, jangan langsung ludes ya."),
  (83,"Gajian!","Catat pemasukan kategori gaji.","Saldo hijau lagi! Tarik napas, ini cuma titipan tagihan."),
  (84,"Tekor Awal Bulan","Catat pengeluaran besar (>=Rp500rb) di tanggal 1-5.","Baru awal bulan udah ngebut. Hati-hati, finish line jauh."),
  (85,"Anak Padang","Catat transaksi mengandung kata 'padang' sebanyak 5 kali.","Rendang lover sejati. Lauknya boleh, dompetnya dijaga."),
  (86,"Patungan Pro","Pakai fitur bagi/split bill (mis. 'makan 200k bagi 4').","Bayar bareng, hemat bareng. Temen-temen sayang kamu."),
  (87,"Tukang Geser","Lakukan 10 kali pindah uang antar saku.","Geser sana geser sini. Bendahara grup ya?"),
  (88,"Nabung Tengah Malam","Lakukan simpan uang antara jam 00.00-04.00.","Insaf tengah malam, langsung nabung. Hidayah finansial."),
  (89,"Begadang Finansial","Catat transaksi antara jam 00.00-04.00.","Mata panda, tapi dompet tetap tercatat. Respect."),
  (90,"Voice Note Master","Pakai input suara sebanyak 10 kali.","Ngomong doang, langsung tercatat. Generasi rebahan."),
  (91,"Ketik Kilat","Catat memakai singkatan nominal (jt/k/rb).","Ngetik 5jt bukan 5000000. Time is money, literally."),
  (92,"Si Telat Catat","Catat transaksi untuk tanggal yang sudah lewat (3 hari atau lebih).","Telat tapi tetap dicatat. Mending telat daripada lupa."),
  (93,"Tukang Edit","Edit transaksi sebanyak 10 kali.","Perfeksionis dompet. Harus pas sampai koma terakhir."),
  (94,"Pembaca Setia","Buka Buku Panduan di tab Profil.","Baca dulu sebelum nanya. Kamu user idaman."),
  (95,"Kepo Fitur","Buka keempat tab: Beranda, Rekapan, Saku, Profil.","Diubek-ubek semua fiturnya. Rasa penasaran tingkat dewa."),
  (96,"Ganti Wajah","Ubah foto profil di tab Profil.","Tampil beda. Dompet rapi, profil juga harus kece."),
  (97,"Kolektor Saku","Miliki 8 saku berbeda.","Dompet bercabang ke mana-mana. Sultan multi-rekening."),
  (98,"Si Kreatif","Buat satu kategori kustom buatan sendiri.","Kategori bawaan kurang? Bikin sendiri, bos. Merdeka!"),
  (99,"Komplit Sehari","Catat pemasukan, pengeluaran, dan pindah uang dalam satu hari yang sama.","Triple combo dalam sehari. Aktivitas dompet padat merayap."),
  (100,"Khatam SakuKilat","Buka semua tab, buka Buku Panduan, dan raih 50 lencana lainnya.","Tamat sudah! Kamu menguasai SakuKilat luar dalam."),
 ]),
]

total = 0
for group_title, items in GROUPS:
    doc.add_heading(group_title, level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["No", "Nama Lencana", "Cara Mendapatkan (Pasti)", "Teks Perayaan"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for no, nama, cara, dop in items:
        total += 1
        row = table.add_row().cells
        row[0].text = str(no)
        row[1].text = nama
        row[2].text = cara
        row[3].text = dop
    doc.add_paragraph("")

doc.add_heading("Catatan", level=1).runs[0].font.color.rgb = CYAN
doc.add_paragraph("Sebagian lencana disiplin/psikologi (mis. 'Hemat Trilogi', 'Survivor Tanggal Tua') dievaluasi "
                  "berdasarkan riwayat bulanan/harianmu dan terbuka saat polanya terpenuhi. Lencana berbasis "
                  "jumlah & data langsung terbuka begitu syarat tercapai, termasuk setelah impor data lama.")
f = doc.add_paragraph(f"Total lencana: {total}. SakuKilat — dibuat oleh Ardhika Argha.")
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
f.runs[0].font.size = Pt(9); f.runs[0].font.color.rgb = GRAY

out = "docs/Panduan-Lengkap-SakuKilat.docx"
doc.save(out)
print(f"SAVED {out} | total achievements = {total}")
